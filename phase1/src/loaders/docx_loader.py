# /rag_data_dispose/phase1/src/loaders/docx_loader.py
from io import BytesIO
from typing import List, Dict, Any
from docx import Document
from .base import DocumentLoader

class DocxLoader(DocumentLoader):
    async def load(self, content: bytes) -> List[Dict[str, Any]]:
        """
        解析 Word (.docx) 文件：
        1. 提取所有段落
        2. 过滤空段落
        3. 每个段落作为一个 block，保留段落样式信息（如标题）
        """
        blocks = []
        doc = Document(BytesIO(content))
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 检测是否为标题（style 包含 Heading）
            is_heading = "Heading" in para.style.name if para.style else False
            
            blocks.append({
                "page_num": 1,
                "text": text,
                "metadata": {
                    "block_index": idx,
                    "style": para.style.name if para.style else "Normal",
                    "is_heading": is_heading,
                    "char_count": len(text)
                }
            })
        
        # 如果没有段落，尝试提取表格内容
        if not blocks and doc.tables:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        blocks.append({
                            "page_num": 1,
                            "text": f"表格第{table_idx+1}行{row_idx+1}: " + row_text,
                            "metadata": {
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "source": "table"
                            }
                        })
        
        return blocks, {}